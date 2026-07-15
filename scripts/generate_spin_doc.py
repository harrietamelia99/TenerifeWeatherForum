import sys
sys.path.insert(0, "/tmp/pydeps")

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colours ──────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0c, 0x23, 0x40)   # #0c2340
GOLD   = RGBColor(0xfb, 0xbf, 0x24)   # #fbbf24
WHITE  = RGBColor(0xff, 0xff, 0xff)
SLATE  = RGBColor(0x47, 0x55, 0x69)   # #475569
LIGHT  = RGBColor(0xf8, 0xfa, 0xfc)   # #f8fafc

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper: set paragraph shading ─────────────────────────────────────────────
def shade_paragraph(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "start", "bottom", "end"):
        tag = OxmlElement(f"w:{ edge }")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        tag.set(qn("w:color"), kwargs.get("color", "auto"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── HEADER BANNER ──────────────────────────────────────────────────────────────
header_para = doc.add_paragraph()
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(header_para, "0c2340")
header_para.paragraph_format.space_before = Pt(10)
header_para.paragraph_format.space_after  = Pt(4)

run = header_para.add_run("Tenerife Weather Forum")
run.font.size  = Pt(10)
run.font.color.rgb = RGBColor(0xfb, 0xbf, 0x24)
run.font.bold  = True

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(title_para, "0c2340")
title_para.paragraph_format.space_before = Pt(0)
title_para.paragraph_format.space_after  = Pt(10)

run2 = title_para.add_run("Lucky Spin Wheel  ✦  How It Works")
run2.font.size  = Pt(22)
run2.font.bold  = True
run2.font.color.rgb = WHITE

doc.add_paragraph()  # spacer

# ── Helper: section heading ────────────────────────────────────────────────────
def add_heading(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    shade_paragraph(p, "1e3a5f")
    run = p.add_run(f"  {text}")
    run.font.bold  = True
    run.font.size  = Pt(13)
    run.font.color.rgb = GOLD

# ── Helper: body paragraph ─────────────────────────────────────────────────────
def add_body(text: str, bold_phrases: list[str] = []):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(6)
    if not bold_phrases:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = SLATE
    else:
        remaining = text
        for phrase in bold_phrases:
            parts = remaining.split(phrase, 1)
            if len(parts) == 2:
                r1 = p.add_run(parts[0])
                r1.font.size = Pt(10.5); r1.font.color.rgb = SLATE
                r2 = p.add_run(phrase)
                r2.font.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = NAVY
                remaining = parts[1]
        r_end = p.add_run(remaining)
        r_end.font.size = Pt(10.5); r_end.font.color.rgb = SLATE

# ── Helper: bullet ─────────────────────────────────────────────────────────────
def add_bullet(text: str, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8)
    if bold_prefix and text.startswith(bold_prefix):
        rb = p.add_run(bold_prefix)
        rb.font.bold = True; rb.font.size = Pt(10.5); rb.font.color.rgb = NAVY
        rr = p.add_run(text[len(bold_prefix):])
        rr.font.size = Pt(10.5); rr.font.color.rgb = SLATE
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5); r.font.color.rgb = SLATE

# ── Helper: highlighted callout box (single-cell table) ────────────────────────
def add_callout(text: str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade_cell(cell, "fef3c7")
    set_cell_border(cell, val="single", sz="6", color="fbbf24")
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x92, 0x40, 0x0e)
    doc.add_paragraph()  # spacer after callout

# ═══════════════════════════════════════════════════════════════════════════════
# CONTENT
# ═══════════════════════════════════════════════════════════════════════════════

# ── 1. The Game ───────────────────────────────────────────────────────────────
add_heading("The Game")
add_body(
    "Every registered user gets one free spin per 24 hours. When they hit the spin button, "
    "the wheel animates and lands on one of 12 segments — each one awarding a different number "
    "of points. Points accumulate across the month and feed into a live leaderboard. At the end "
    "of each month, the top 3 players win a prize, the leaderboard resets to zero, and a new "
    "competition begins.",
    bold_phrases=["one free spin per 24 hours", "top 3 players win a prize"]
)

# ── 2. Wheel Segments ─────────────────────────────────────────────────────────
add_heading("The Wheel Segments")
add_body(
    "The wheel has 12 segments, all Tenerife-themed. Not all segments are equally likely — "
    "the rarer, higher-value prizes come up less often, keeping the game exciting but fair.",
)

segments = [
    ("Grand Prize",   "500 pts", "Very rare"),
    ("Jackpot",       "300 pts", "Rare"),
    ("Hot Streak",    "200 pts", "Uncommon"),
    ("Sunny Day",     "150 pts", "Uncommon"),
    ("Bonus Points",  "100 pts", "Moderate"),
    ("Lucky Dip",     "75 pts",  "Moderate"),
    ("Warm Breeze",   "50 pts",  "Common"),
    ("Trade Winds",   "40 pts",  "Common"),
    ("Sun Shower",    "30 pts",  "Common"),
    ("Cloudy Day",    "20 pts",  "Very common"),
    ("Mist & Fog",    "10 pts",  "Very common"),
    ("Spin Again",    "0 pts",   "Occasional — free re-spin, no daily allowance used"),
]

tbl = doc.add_table(rows=len(segments) + 1, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Segment", "Points", "Likelihood"]
for i, h in enumerate(headers):
    cell = tbl.cell(0, i)
    shade_cell(cell, "0c2340")
    set_cell_border(cell, val="single", sz="4", color="fbbf24")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = GOLD

for row_idx, (name, pts, likelihood) in enumerate(segments):
    fill = "f8fafc" if row_idx % 2 == 0 else "ffffff"
    row = tbl.rows[row_idx + 1]
    for col_idx, val in enumerate([name, pts, likelihood]):
        cell = row.cells[col_idx]
        shade_cell(cell, fill)
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY if col_idx == 0 else SLATE
        if col_idx == 0: r.font.bold = True

doc.add_paragraph()

# ── 3. Daily Spin Limit ───────────────────────────────────────────────────────
add_heading("How the Daily Spin Limit Works")
add_body(
    "The spin limit is enforced server-side, not just in the browser. When a user clicks spin, "
    "the server checks their last spin timestamp in the database. If it was less than 24 hours "
    "ago, the request is rejected — there is no way to bypass this by refreshing the page or "
    "opening multiple tabs.",
    bold_phrases=["server-side"]
)
add_body(
    "To prevent someone clicking the button twice quickly and getting two spins, the system uses "
    "an atomic database update — meaning the spin slot is claimed in a single operation that "
    "only succeeds once. A countdown timer on the page shows the user exactly how long until "
    "their next spin, down to the second.",
    bold_phrases=["atomic database update"]
)

# ── 4. The Winning Result ─────────────────────────────────────────────────────
add_heading("The Winning Result")
add_body(
    "The winning segment is chosen server-side before the animation even starts. The server picks "
    "the result using weighted probabilities, sends it to the browser, and the wheel animation "
    "is then choreographed to land precisely on that segment. The result can never be tampered "
    "with from the browser — what the server decides is final. Once the wheel stops, a full-screen "
    "celebration pop-up shows the user what they won.",
    bold_phrases=["server-side", "final"]
)

# ── 5. Points & Leaderboard ───────────────────────────────────────────────────
add_heading("Points & the Leaderboard")
add_body("Every user has two points totals:")
add_bullet("Monthly points — resets to zero each competition month. This is what drives the leaderboard.", "Monthly points")
add_bullet("Lifetime points — a running total that never resets. Personal achievement tracking.", "Lifetime points")
add_body(
    "The leaderboard shows the current top 10 players. If two players are tied on the same "
    "monthly score, the one who reached that score first ranks higher — rewarding consistency "
    "and early play rather than a last-minute rush.",
    bold_phrases=["reached that score first"]
)

# ── 6. Newsletter Bonus Spin ──────────────────────────────────────────────────
add_heading("Newsletter Bonus Spin")
add_body(
    "Newsletter subscribers can be rewarded with a bonus spin via the admin panel. The bonus "
    "spin can be used within the 24-hour cooldown window, giving subscribers an extra go even "
    "if they have already spun that day. It can only be used once before it needs to be re-granted — "
    "so it is a meaningful, controllable reward rather than an unlimited loophole.",
    bold_phrases=["bonus spin"]
)

# ── 7. Monthly Reset & Emails ─────────────────────────────────────────────────
add_heading("Monthly Reset & Winner Emails")
add_body("At the end of each month, you click Archive & Reset Month in the admin panel. This automatically:")
add_bullet("Saves the top 3 players and their scores permanently to a historical archive")
add_bullet("Sends a branded congratulations email to each of the top 3 winners")
add_bullet("Sends you a summary email listing all three winners' names, emails and scores")
add_bullet("Resets every player's monthly points to zero so the new month starts fresh")
add_callout(
    "Winner emails go out automatically — you just need to follow up with the actual prizes. "
    "The summary email lands in your inbox with everything you need to get in touch."
)

# ── 8. Admin Panel ────────────────────────────────────────────────────────────
add_heading("The Admin Panel")
add_body(
    "The admin panel is password-protected and not linked from anywhere on the public site. "
    "You can access it here:",
    bold_phrases=[]
)
p_link = doc.add_paragraph()
p_link.paragraph_format.space_before = Pt(2)
p_link.paragraph_format.space_after  = Pt(8)
r_link = p_link.add_run("https://www.tenerifeweatherforum.com/preview/spin/admin")
r_link.font.bold  = True
r_link.font.size  = Pt(10.5)
r_link.font.color.rgb = NAVY

add_body("When you visit that page you will be prompted for your admin password (sent separately). Inside you will find:")
add_bullet("Leaderboard tab — current month's top players ranked by points")
add_bullet("Users tab — every registered account with monthly points, lifetime points, last spin time and bonus spin status")
add_bullet("Adjust Points — manually increase or decrease any user's points if a correction is ever needed")
add_bullet("Grant Bonus Spin — give a specific user an extra spin, ideal for rewarding newsletter subscribers")
add_bullet("Archive & Reset Month — closes the month, fires winner emails, and resets all scores to zero")

# ── 9. Account System ─────────────────────────────────────────────────────────
add_heading("Account System")
add_body(
    "Users register with an email address, a display name shown publicly on the leaderboard, "
    "and a password. Passwords are encrypted using industry-standard hashing before storage — "
    "they are never saved in plain text anywhere. Sessions are handled with signed tokens that "
    "expire automatically. The spin registration pages are completely separate from the main site.",
    bold_phrases=["never saved in plain text"]
)

# ── 10. Security Summary ──────────────────────────────────────────────────────
add_heading("Security & Fairness Summary")
add_bullet("All spin results are decided on the server — the browser cannot influence the outcome")
add_bullet("The 24-hour cooldown is enforced at database level and cannot be bypassed client-side")
add_bullet("Atomic database operations prevent double-spins from rapid clicking or multiple tabs")
add_bullet("Passwords are hashed before storage — never stored in plain text")
add_bullet("The admin panel requires a separate password and is not linked from the public site")

# ── 11. Current Status ────────────────────────────────────────────────────────
add_heading("Current Status — Awaiting Your Approval")
add_callout(
    "The wheel is currently live in test mode at a hidden preview link. "
    "Nothing is visible to the public yet — it will only go live once you give the go-ahead."
)
add_body("You can view and test the full game here:")
p_spin = doc.add_paragraph()
p_spin.paragraph_format.space_before = Pt(2)
p_spin.paragraph_format.space_after  = Pt(8)
r_spin = p_spin.add_run("https://www.tenerifeweatherforum.com/preview/spin")
r_spin.font.bold  = True
r_spin.font.size  = Pt(10.5)
r_spin.font.color.rgb = NAVY

add_body(
    "Simply create a free account on that page to register and start spinning. Everything works "
    "exactly as it will when it goes live — the spin logic, the leaderboard, the cooldown timer, "
    "winner emails and the admin panel are all fully functional.",
)
add_body(
    "Once you are happy with how it looks and works, just say the word and we will add it to "
    "the main site navigation. We will also add a section to the homepage at that point to "
    "introduce the competition and encourage visitors to sign up.",
    bold_phrases=["just say the word"]
)

# ── FOOTER ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(footer_p, "0c2340")
footer_p.paragraph_format.space_before = Pt(8)
footer_p.paragraph_format.space_after  = Pt(8)
rf = footer_p.add_run("Tenerife Weather Forum  ·  tenerifeweatherforum.com")
rf.font.size = Pt(9)
rf.font.color.rgb = RGBColor(0xfb, 0xbf, 0x24)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/harrietpearce/Tenerife Weather Forum/Lucky Spin Wheel — How It Works.docx"
doc.save(out)
print(f"Saved: {out}")
